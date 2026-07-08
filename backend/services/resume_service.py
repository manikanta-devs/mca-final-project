import logging
import re
import os
from services.resume_keywords import SKILL_KEYWORDS, EDUCATION_KEYWORDS, EXPERIENCE_PATTERNS

logger = logging.getLogger(__name__)


class ResumeService:
    """Service for extracting and analyzing resume content"""

    SKILL_KEYWORDS = SKILL_KEYWORDS
    EDUCATION_KEYWORDS = EDUCATION_KEYWORDS
    EXPERIENCE_PATTERNS = EXPERIENCE_PATTERNS

    def __init__(self):
        self.nlp = self._load_spacy()

    def _load_spacy(self):
        """Load spaCy model with fallback"""
        try:
            import spacy

            try:
                return spacy.load("en_core_web_sm")
            except OSError:
                banner = (
                    "\n"
                    "================================================================================\n"
                    "⚠️  WARNING: spaCy model 'en_core_web_sm' was not found!\n"
                    "   Resume parsing accuracy will be degraded (falling back to regex matches).\n"
                    "   Fix this by downloading the model: python -m spacy download en_core_web_sm\n"
                    "================================================================================\n"
                )
                print(banner)
                logger.warning(
                    "spaCy model not found. Run: python -m spacy download en_core_web_sm"
                )
                return None
        except Exception as exc:
            banner = (
                "\n"
                "================================================================================\n"
                "⚠️  WARNING: spaCy library is not available!\n"
                "   Resume parsing will use basic regex fallback routines.\n"
                f"   Error: {exc}\n"
                "================================================================================\n"
            )
            print(banner)
            logger.warning(
                "spaCy unavailable, falling back to regex extraction: %s", exc
            )
            return None

    def extract_text_from_pdf(self, filepath: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2

            text = []
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return "\n".join(text)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to extract PDF text: {e}")

    def extract_text_from_docx(self, filepath: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document

            doc = Document(filepath)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            # Extract from tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            return "\n".join(text)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Failed to extract DOCX text: {e}")

    def extract_text_from_txt(self, filepath: str) -> str:
        """Extract text from TXT file"""
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def analyze_resume(self, filepath: str) -> dict:
        """Main method to analyze a resume file"""
        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".pdf":
            text = self.extract_text_from_pdf(filepath)
        elif ext in [".docx", ".doc"]:
            text = self.extract_text_from_docx(filepath)
        elif ext == ".txt":
            text = self.extract_text_from_txt(filepath)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        return self.analyze_text(text)

    def analyze_text(self, text: str) -> dict:
        """Analyze resume text and extract structured information"""
        text_lower = text.lower()

        skills = self._extract_skills(text_lower)
        education = self._extract_education(text, text_lower)
        experience = self._extract_experience(text, text_lower)
        contact = self._extract_contact(text)
        entities = self._extract_entities(text)
        score = self._calculate_resume_score(skills, education, experience, text)

        def clean_json(raw_text: str) -> str:
            if not raw_text:
                return ""
            t = raw_text.strip()
            first_brace = t.find('{')
            last_brace = t.rfind('}')
            if first_brace != -1 and last_brace != -1:
                return t[first_brace:last_brace+1]
            return t

        deep_analysis = None
        try:
            from ai.gemini_service import GeminiService
            gemini = GeminiService()
            prompt = f"""You are an ATS expert and career coach. Analyze this resume and return ONLY a JSON object. Be extremely concise — keep every string under 15 words. Limit all arrays to max 3 items.
Resume:
---
{text[:3000]}
---
Return exactly this JSON structure (no markdown, no extra text), replacing the placeholders (like score_0_to_100, status_excellent_or_needs_improvement_or_weak_or_missing) with actual calculated, realistic values based on the resume:
{{
  "coach_report": {{
    "summary": "1-sentence summary of the candidate's profile",
    "strengths": ["strength1", "strength2", "strength3"],
    "weaknesses": ["weakness1", "weakness2", "weakness3"],
    "missing_keywords": ["kw1", "kw2", "kw3"],
    "missing_sections": ["sec1", "sec2"],
    "grammar_suggestions": ["suggestion1"],
    "actionable_improvements": ["action1", "action2", "action3"],
    "current_score": current_score_0_to_100,
    "potential_score": potential_score_0_to_100,
    "next_steps": ["step1", "step2"]
  }},
  "heatmap": {{
    "name": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "contact": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "summary": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "skills": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "experience": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "projects": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "education": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "certifications": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "achievements": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "keywords": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}},
    "formatting": {{"status": "status_excellent_or_needs_improvement_or_weak_or_missing", "score": score_0_to_100, "feedback": "Brief feedback."}}
  }},
  "interview_prep": {{
    "estimated_duration": "30 Minutes",
    "company_styles": ["Google", "Amazon", "Microsoft"],
    "questions": [
      {{"text": "Tailored question text?", "category": "technical", "difficulty": "medium", "common_mistakes": ["mistake1"]}},
      {{"text": "Tailored question text?", "category": "behavioral", "difficulty": "easy", "common_mistakes": ["mistake1"]}},
      {{"text": "Tailored question text?", "category": "project", "difficulty": "hard", "common_mistakes": ["mistake1"]}}
    ]
  }},
  "career_roadmap": {{
    "current_level": "Beginner_or_Intermediate_or_Advanced",
    "suitable_roles": ["Role1", "Role2", "Role3"],
    "salary_range": "Salary range estimate",
    "missing_skills": ["skill1", "skill2"],
    "learning_path": [
      {{"week": "Week 1", "topic": "Topic name", "detail": "Short detail."}},
      {{"week": "Week 2", "topic": "Topic name", "detail": "Short detail."}},
      {{"week": "Week 3", "topic": "Topic name", "detail": "Short detail."}}
    ],
    "recommended_certifications": ["Cert1", "Cert2"],
    "job_readiness_percentage": readiness_percentage_0_to_100,
    "estimated_time": "Months_estimate"
  }}
}}"""
            raw = gemini.generate_content(prompt, max_tokens=1800)
            if raw:
                import json
                cleaned = clean_json(raw)
                deep_analysis = json.loads(cleaned)
        except Exception as e:
            logger.error(f"Deep AI analysis failed: {e}")

        if not deep_analysis:
            # Build a dynamic fallback using the ACTUAL extracted resume data so
            # each resume shows different values even when AI providers are on cooldown.
            _has_skills    = bool(skills.get("all"))
            _has_exp       = bool(experience.get("years") or experience.get("titles"))
            _has_edu       = bool(education)
            _has_contact   = bool(contact.get("email") or contact.get("phone"))
            _has_github    = bool(contact.get("github"))
            _skill_list    = skills.get("all", [])[:6]
            _exp_level     = experience.get("level", "entry")
            _exp_years     = experience.get("years")
            _top_edu       = education[0] if education else "Degree not found"
            _base_score    = score.get("percentage", 60) if isinstance(score, dict) else int(score or 60)

            # Heatmap scores derived from actual extracted fields
            _exp_score = 80 if _exp_years and _exp_years >= 3 else (60 if _has_exp else 35)
            _skills_score = min(95, 50 + len(_skill_list) * 7)
            _edu_score = 90 if _has_edu else 45
            _contact_score = 95 if _has_contact else 55
            _github_score = 85 if _has_github else 40

            _missing_kws = [k for k in ["Docker", "AWS", "REST APIs", "CI/CD", "Unit Testing", "Kubernetes"]
                            if k.lower() not in " ".join(_skill_list).lower()][:4]
            _missing_secs = []
            if not _has_exp:      _missing_secs.append("Work Experience")
            if not _has_github:   _missing_secs.append("GitHub Portfolio")
            if len(_skill_list) < 4: _missing_secs.append("Skills Section")

            _roadmap_roles = {
                "senior": ["Senior Software Engineer", "Tech Lead", "Solutions Architect"],
                "mid":    ["Software Engineer", "Full Stack Developer", "Backend Developer"],
                "junior": ["Junior Developer", "Associate Engineer", "Frontend Developer"],
                "entry":  ["Junior Developer", "Trainee Engineer", "Intern"],
            }.get(_exp_level, ["Software Engineer", "Developer", "Engineer"])

            _readiness = min(95, max(40, _base_score))

            deep_analysis = {
                "coach_report": {
                    "summary": f"Resume shows {'strong' if _base_score >= 75 else 'moderate'} technical profile with {len(_skill_list)} detected skills and {'solid' if _has_exp else 'limited'} work experience.",
                    "strengths": (
                        ([f"Strong technical skills: {', '.join(_skill_list[:3])}"] if _skill_list else []) +
                        (["Verified education credentials"] if _has_edu else []) +
                        (["GitHub portfolio present"] if _has_github else ["Clean contact information"] if _has_contact else [])
                    )[:3] or ["Technical foundation present"],
                    "weaknesses": (
                        ([] if _has_exp else ["No professional work experience listed"]) +
                        ([] if _has_github else ["No GitHub or portfolio link"]) +
                        (["Missing key DevOps/cloud skills"] if _missing_kws else [])
                    )[:3] or ["Consider adding quantified achievements"],
                    "missing_keywords": _missing_kws,
                    "missing_sections": _missing_secs or ["Certifications"],
                    "grammar_suggestions": ["Use active verbs in all project and experience bullet points."],
                    "actionable_improvements": [
                        "Add measurable impact metrics to each project.",
                        "Link your GitHub profile and highlight top repositories.",
                        "Obtain at least one cloud certification (AWS/GCP).",
                    ],
                    "current_score": _base_score,
                    "potential_score": min(98, _base_score + 18),
                    "next_steps": [
                        "Start a mock interview tailored to your detected skills.",
                        "Add quantitative results to your project descriptions.",
                    ]
                },
                "heatmap": {
                    "name":           {"status": "excellent",          "score": 95,           "feedback": "Name is clearly visible and correctly formatted.", "recruiter_view": "Perfect first impression.", "ats_view": "Name parsed cleanly.", "suggested_rewrite": "N/A"},
                    "contact":        {"status": "excellent" if _contact_score >= 80 else "needs_improvement", "score": _contact_score, "feedback": "Contact info present." if _has_contact else "Email or phone missing.", "recruiter_view": "Easy to reach." if _has_contact else "Hard to contact candidate.", "ats_view": "Contact fields indexed." if _has_contact else "No contact fields found.", "suggested_rewrite": "N/A"},
                    "summary":        {"status": "needs_improvement",  "score": 65,           "feedback": "Summary section needs more targeted keywords and impact statements.", "recruiter_view": "Reads as generic.", "ats_view": "Lacks high-impact job keywords.", "suggested_rewrite": "Results-driven engineer with proven track record in " + (", ".join(_skill_list[:2]) if _skill_list else "software development") + "."},
                    "skills":         {"status": "excellent" if _skills_score >= 75 else "needs_improvement", "score": _skills_score, "feedback": f"{len(_skill_list)} skills detected: {', '.join(_skill_list[:3])}." if _skill_list else "No skills section found.", "recruiter_view": "Good technical stack." if _skill_list else "Skills section empty.", "ats_view": "Keyword density adequate." if _skill_list else "No skill keywords.", "suggested_rewrite": "N/A"},
                    "experience":     {"status": "excellent" if _exp_score >= 75 else ("needs_improvement" if _exp_score >= 50 else "weak_or_missing"), "score": _exp_score, "feedback": f"{_exp_years} years experience detected." if _exp_years else ("Job titles found but no duration." if _has_exp else "No work experience listed."), "recruiter_view": "Adequate experience." if _has_exp else "No work history — flag for screening.", "ats_view": "Experience duration parsed." if _exp_years else "No experience duration found.", "suggested_rewrite": "Add freelance, internship, or part-time roles with dates." if not _has_exp else "N/A"},
                    "projects":       {"status": "needs_improvement",  "score": 62,           "feedback": "Projects lack measurable outcomes and impact metrics.", "recruiter_view": "Hard to gauge complexity without numbers.", "ats_view": "Project keywords present but no metrics.", "suggested_rewrite": "Reduced page load time by 40% using lazy loading and code splitting."},
                    "education":      {"status": "excellent" if _edu_score >= 80 else "needs_improvement", "score": _edu_score, "feedback": _top_edu if _has_edu else "Education section not detected.", "recruiter_view": "Meets requirements." if _has_edu else "No education found.", "ats_view": "Education indexed." if _has_edu else "Missing education keywords.", "suggested_rewrite": "N/A"},
                    "certifications": {"status": "weak_or_missing",    "score": 35,           "feedback": "No professional certifications found.", "recruiter_view": "Add cloud or tech certs.", "ats_view": "Zero certification keywords.", "suggested_rewrite": "AWS Certified Cloud Practitioner (2025)"},
                    "achievements":   {"status": "weak_or_missing",    "score": 35,           "feedback": "No dedicated achievements or awards section.", "recruiter_view": "Misses extracurricular signals.", "ats_view": "No achievement keywords found.", "suggested_rewrite": "Top 10% in university batch or hackathon finalist."},
                    "keywords":       {"status": "needs_improvement" if _missing_kws else "excellent", "score": 60 if _missing_kws else 85, "feedback": f"Missing: {', '.join(_missing_kws)}." if _missing_kws else "Good keyword coverage.", "recruiter_view": "Needs more job-specific terms." if _missing_kws else "Strong keyword match.", "ats_view": "Some JD keywords missing." if _missing_kws else "Good ATS keyword density.", "suggested_rewrite": "N/A"},
                    "formatting":     {"status": "excellent",          "score": 85,           "feedback": "Clean structure and easy to parse.", "recruiter_view": "Scans well visually.", "ats_view": "Layout is ATS-friendly.", "suggested_rewrite": "N/A"},
                },
                "interview_prep": {
                    "estimated_duration": "30 Minutes",
                    "company_styles": ["Google", "Amazon", "Microsoft"],
                    "questions": [
                        {
                            "text": f"Walk me through your experience with {_skill_list[0] if _skill_list else 'your primary technology'} and give a real project example.",
                            "category": "technical",
                            "difficulty": "medium",
                            "expected_answer": "Use STAR method. Mention specific libraries, challenges, and outcomes.",
                            "key_concepts": _skill_list[:3] or ["Core concepts"],
                            "common_mistakes": ["Being too vague without a concrete example."]
                        },
                        {
                            "text": "Tell me about a project where you had to solve a hard technical problem. What was your approach?",
                            "category": "project",
                            "difficulty": "medium",
                            "expected_answer": "Describe problem identification, solution design, implementation, and measurable result.",
                            "key_concepts": ["Problem solving", "System design", "Trade-offs"],
                            "common_mistakes": ["Focusing on the task, not the impact or result."]
                        },
                        {
                            "text": "How do you handle disagreement with a team member or tech lead?",
                            "category": "behavioral",
                            "difficulty": "easy",
                            "expected_answer": "Show collaboration, active listening, and data-driven resolution.",
                            "key_concepts": ["Communication", "Collaboration", "Objectivity"],
                            "common_mistakes": ["Sounding defensive or avoiding conflict detail."]
                        }
                    ]
                },
                "career_roadmap": {
                    "current_level": _exp_level.capitalize(),
                    "suitable_roles": _roadmap_roles,
                    "salary_range": "INR 4-10 LPA" if _exp_level in ["entry", "junior"] else "INR 10-25 LPA",
                    "missing_skills": _missing_kws or ["Cloud Platforms", "System Design"],
                    "learning_path": [
                        {"week": "Week 1", "topic": _missing_kws[0] if _missing_kws else "System Design", "detail": f"Master {'containers and deployment' if 'Docker' in (_missing_kws or []) else 'fundamentals and best practices'}."},
                        {"week": "Week 2", "topic": _missing_kws[1] if len(_missing_kws) > 1 else "Cloud Fundamentals", "detail": "Build hands-on projects and push to GitHub."},
                        {"week": "Week 3", "topic": "Mock Interviews", "detail": "Practice 3 full mock interviews focusing on your weak areas."},
                    ],
                    "recommended_certifications": ["AWS Certified Cloud Practitioner", "Google Associate Cloud Engineer"],
                    "job_readiness_percentage": _readiness,
                    "estimated_time": "2 Months" if _base_score >= 70 else "3 Months"
                }
            }

        res = {
            "raw_text_length": len(text),
            "skills": skills,
            "education": education,
            "experience": experience,
            "contact": contact,
            "entities": entities,
            "resume_score": score,
            "summary": self._generate_summary(skills, education, experience),
        }
        res.update(deep_analysis)
        return res

    def _extract_skills(self, text_lower: str) -> dict:
        """Extract skills by category"""
        found = {}
        all_skills = []
        for category, keywords in self.SKILL_KEYWORDS.items():
            matched = []
            for kw in keywords:
                if kw in text_lower:
                    matched.append(kw.title() if len(kw) > 3 else kw.upper())
            if matched:
                found[category] = list(set(matched))
                all_skills.extend(matched)
        found["all"] = list(set(all_skills))
        return found

    def _extract_education(self, text: str, text_lower: str) -> list:
        """Extract education information"""
        education = []
        lines = text.split("\n")
        for line in lines:
            line_lower = line.lower()
            if any(kw in line_lower for kw in self.EDUCATION_KEYWORDS):
                clean = line.strip()
                if len(clean) > 10 and clean not in education:
                    education.append(clean)
        return education[:5]  # Limit to 5 entries

    def _extract_experience(self, text: str, text_lower: str) -> dict:
        """Extract work experience information"""
        years = None
        for pattern in self.EXPERIENCE_PATTERNS:
            match = re.search(pattern, text_lower)
            if match:
                years = int(match.group(1))
                break

        job_titles = []
        title_keywords = [
            "engineer",
            "developer",
            "manager",
            "analyst",
            "designer",
            "architect",
            "lead",
            "senior",
            "junior",
            "intern",
            "consultant",
            "specialist",
            "director",
            "head of",
            "vp",
            "cto",
            "ceo",
        ]
        lines = text.split("\n")
        for line in lines:
            line_lower = line.lower()
            if any(kw in line_lower for kw in title_keywords):
                clean = line.strip()
                if 5 < len(clean) < 80 and clean not in job_titles:
                    job_titles.append(clean)

        level = "entry"
        if years:
            if years >= 8:
                level = "senior"
            elif years >= 4:
                level = "mid"
            elif years >= 1:
                level = "junior"

        return {"years": years, "level": level, "titles": job_titles[:5]}

    def _extract_contact(self, text: str) -> dict:
        """Extract contact information"""
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        phone_pattern = r"[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}"
        linkedin_pattern = r"linkedin\.com/in/[\w-]+"
        github_pattern = r"github\.com/[\w-]+"

        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        linkedin = re.findall(linkedin_pattern, text.lower())
        github = re.findall(github_pattern, text.lower())

        return {
            "email": emails[0] if emails else None,
            "phone": phones[0] if phones else None,
            "linkedin": linkedin[0] if linkedin else None,
            "github": github[0] if github else None,
        }

    def _extract_entities(self, text: str) -> dict:
        """Use spaCy to extract named entities"""
        if not self.nlp:
            return {"organizations": [], "locations": [], "persons": []}

        try:
            doc = self.nlp(text[:5000])  # Limit for performance
            entities = {"organizations": [], "locations": [], "persons": []}

            for ent in doc.ents:
                if ent.label_ == "ORG" and ent.text not in entities["organizations"]:
                    entities["organizations"].append(ent.text)
                elif (
                    ent.label_ in ["GPE", "LOC"]
                    and ent.text not in entities["locations"]
                ):
                    entities["locations"].append(ent.text)
                elif ent.label_ == "PERSON" and ent.text not in entities["persons"]:
                    entities["persons"].append(ent.text)

            return {k: v[:10] for k, v in entities.items()}
        except Exception as e:
            logger.error(f"spaCy entity extraction error: {e}")
            return {"organizations": [], "locations": [], "persons": []}

    def _calculate_resume_score(
        self, skills: dict, education: list, experience: dict, text: str
    ) -> dict:
        """Calculate a resume quality score"""
        scores = {}
        # Skills score (0-25)
        skill_count = len(skills.get("all", []))
        scores["skills"] = min(25, skill_count * 2)

        # Education score (0-25)
        scores["education"] = min(25, len(education) * 8)

        # Experience score (0-25)
        exp_years = experience.get("years", 0) or 0
        scores["experience"] = min(25, exp_years * 4)

        # Completeness score (0-25)
        completeness = 0
        if len(text) > 500:
            completeness += 5
        if len(text) > 1000:
            completeness += 5
        if "email" in text.lower() or "@" in text:
            completeness += 5
        if any(w in text.lower() for w in ["project", "achievement", "award"]):
            completeness += 5
        if len(education) > 0:
            completeness += 5
        scores["completeness"] = completeness

        total = sum(scores.values())
        grades = {90: "A+", 80: "A", 70: "B+", 60: "B", 50: "C+", 0: "C"}
        grade = next(
            g
            for threshold, g in sorted(grades.items(), reverse=True)
            if total >= threshold
        )

        return {
            "total": total,
            "max": 100,
            "grade": grade,
            "breakdown": scores,
            "percentage": total,
        }

    def _generate_summary(self, skills: dict, education: list, experience: dict) -> str:
        """Generate a brief resume summary"""
        parts = []
        exp_years = experience.get("years")
        level = experience.get("level", "entry")

        if exp_years:
            parts.append(f"{exp_years} years of experience")

        skill_count = len(skills.get("all", []))
        if skill_count > 0:
            top_skills = skills.get("all", [])[:3]
            parts.append(f"skilled in {', '.join(top_skills)}")

        if education:
            parts.append("with formal education background")

        if parts:
            return f"{level.title()}-level professional {', '.join(parts)}."
        return "Resume processed successfully."

    def compare_resume_to_job(self, resume_data: dict, job_description: str) -> dict:
        """Compare analyzed resume data against a target job description."""
        if not resume_data:
            raise ValueError("Resume data is required for job matching")

        if not job_description or len(job_description.strip()) < 30:
            raise ValueError("Job description is required for job matching")

        job_analysis = self.analyze_text(job_description)

        resume_skill_map = self._skill_map(resume_data.get("skills", {}).get("all", []))
        job_skill_map = self._skill_map(job_analysis.get("skills", {}).get("all", []))

        resume_skill_keys = set(resume_skill_map.keys())
        job_skill_keys = set(job_skill_map.keys())

        matched_keys = sorted(resume_skill_keys & job_skill_keys)
        missing_keys = sorted(job_skill_keys - resume_skill_keys)
        extra_keys = sorted(resume_skill_keys - job_skill_keys)

        matched_skills = [
            resume_skill_map.get(key) or job_skill_map.get(key) for key in matched_keys
        ]
        missing_skills = [job_skill_map[key] for key in missing_keys]
        extra_skills = [resume_skill_map[key] for key in extra_keys]

        resume_experience = resume_data.get("experience", {}) or {}
        job_experience = job_analysis.get("experience", {}) or {}
        resume_years = int(resume_experience.get("years") or 0)
        required_years = int(job_experience.get("years") or 0)

        skill_coverage = (
            len(matched_skills) / len(job_skill_keys) if job_skill_keys else 0.65
        )
        experience_alignment = 1.0
        if required_years > 0:
            experience_alignment = (
                min(1.0, resume_years / required_years) if resume_years else 0.0
            )

        resume_completion = (resume_data.get("resume_score", {}) or {}).get(
            "percentage", 0
        ) / 100
        education_alignment = 1.0 if resume_data.get("education") else 0.6

        total_score = round(
            (skill_coverage * 58)
            + (experience_alignment * 20)
            + (education_alignment * 10)
            + (min(1.0, resume_completion) * 12)
        )
        total_score = max(0, min(100, total_score))

        grade = self._grade_from_score(total_score)
        readiness = self._readiness_label(total_score)

        strengths = []
        if matched_skills:
            strengths.append(f"Matches {len(matched_skills)} job keywords")
        if required_years and resume_years >= required_years:
            strengths.append("Experience level meets or exceeds the job signal")
        if resume_data.get("education"):
            strengths.append("Education section is present")
        if resume_data.get("summary"):
            strengths.append("Resume already has a concise summary")

        recommendations = []
        if missing_skills:
            recommendations.append(
                f"Add or emphasize these keywords: {', '.join(missing_skills[:5])}."
            )
        if required_years and resume_years < required_years:
            recommendations.append(
                f"Highlight projects that show at least {required_years - resume_years} more year(s) of relevant impact."
            )
        if not resume_data.get("contact", {}).get("email"):
            recommendations.append(
                "Make sure your contact details are visible and complete."
            )
        if len(recommendations) < 3:
            recommendations.append(
                "Tailor your professional summary to mirror the language in the job description."
            )
            recommendations.append(
                "Quantify outcomes in bullets with metrics, scale, or impact."
            )

        priority_actions = [
            "Rewrite the top summary line with the role title and 2-3 matching keywords.",
            "Add one bullet for each missing priority skill using a concrete project example.",
            "Prepare interview stories around the matched skills and the strongest project wins.",
        ]

        return {
            "match_score": total_score,
            "match_grade": grade,
            "readiness": readiness,
            "matched_skills": matched_skills[:12],
            "missing_skills": missing_skills[:12],
            "extra_skills": extra_skills[:12],
            "resume_years": resume_years,
            "required_years": required_years,
            "skill_coverage": round(skill_coverage * 100, 1),
            "experience_alignment": round(experience_alignment * 100, 1),
            "education_alignment": round(education_alignment * 100, 1),
            "strengths": strengths[:4],
            "recommendations": recommendations[:4],
            "priority_actions": priority_actions,
            "job_keywords": list(job_skill_map.values())[:12],
            "summary": self._build_match_summary(
                total_score,
                matched_skills,
                missing_skills,
                required_years,
                resume_years,
            ),
        }

    def _skill_map(self, skills: list) -> dict:
        """Normalize skill labels for matching while keeping display text."""
        mapped = {}
        for skill in skills or []:
            if not skill:
                continue
            display = skill.strip()
            normalized = display.lower()
            mapped[normalized] = display
        return mapped

    def _grade_from_score(self, score: int) -> str:
        """Convert a numeric score to a letter grade."""
        grade_map = [(90, "A+"), (80, "A"), (70, "B+"), (60, "B"), (50, "C+"), (0, "C")]
        return next(grade for threshold, grade in grade_map if score >= threshold)

    def _readiness_label(self, score: int) -> str:
        """Translate the fit score into a hiring-readiness label."""
        if score >= 85:
            return "interview-ready"
        if score >= 70:
            return "strong-match"
        if score >= 55:
            return "build-gap"
        return "foundation-needed"

    def _build_match_summary(
        self,
        score: int,
        matched_skills: list,
        missing_skills: list,
        required_years: int,
        resume_years: int,
    ) -> str:
        """Generate a short executive summary for the match report."""
        if score >= 85:
            prefix = "Excellent alignment"
        elif score >= 70:
            prefix = "Strong alignment"
        elif score >= 55:
            prefix = "Moderate alignment"
        else:
            prefix = "Low alignment"

        skill_text = f"{len(matched_skills)} matched skills"
        gap_text = f"{len(missing_skills)} keyword gaps"
        if required_years:
            exp_text = f"{resume_years}/{required_years} years of experience signal"
        else:
            exp_text = f"{resume_years} years of experience signal"

        return f"{prefix}: {skill_text}, {gap_text}, {exp_text}."
